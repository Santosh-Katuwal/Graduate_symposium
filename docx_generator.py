"""
docx_generator.py — Generates a formatted .docx report that replicates
the layout defined in Template.tex (headshot + info block, abstract, figures).
"""

import io
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from PIL import Image

import math2docx
import config as C


# ── Helpers ──────────────────────────────────────────────────────────────────

def _apply(run, font_size=None, bold=False, italic=False):
    """Apply consistent font styling to a single run."""
    run.font.name = C.FONT_NAME
    run.font.size = font_size or C.FONT_SIZE_BODY
    run.font.color.rgb = RGBColor.from_string(C.FONT_COLOR_HEX)
    run.bold = bold
    run.italic = italic


def _add_parsed_text(paragraph, text, font_size=None, bold=False, italic=False):
    """Parse text for $$...$$ and $...$ blocks and insert as MathML or regular text."""
    if not text:
        return
    font_size = font_size or C.FONT_SIZE_BODY
    pattern = re.compile(r'(\$\$.*?\$\$|\$.*?\$)', re.DOTALL)
    parts = re.split(pattern, text)
    
    for part in parts:
        if not part:
            continue
        if part.startswith('$$') and part.endswith('$$'):
            math_str = part[2:-2].strip()
            try:
                math2docx.add_math(paragraph, math_str)
            except Exception:
                run = paragraph.add_run(part)
                _apply(run, font_size=font_size, bold=bold, italic=italic)
        elif part.startswith('$') and part.endswith('$'):
            math_str = part[1:-1].strip()
            try:
                math2docx.add_math(paragraph, math_str)
            except Exception:
                run = paragraph.add_run(part)
                _apply(run, font_size=font_size, bold=bold, italic=italic)
        else:
            run = paragraph.add_run(part)
            _apply(run, font_size=font_size, bold=bold, italic=italic)


def _set_cell_margins(cell, top=0, bottom=0, start=0, end=0):
    """Set cell margins in EMUs (1 inch = 914400 EMUs)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn("w:tcMar"))
    if tcMar is None:
        tcMar = tc.makeelement(qn("w:tcMar"), {})
        tcPr.append(tcMar)
    for side, val in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        el = tcMar.find(qn(f"w:{side}"))
        if el is None:
            el = tc.makeelement(qn(f"w:{side}"), {})
            tcMar.append(el)
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")


def _remove_table_borders(table):
    """Remove all borders from a table so it looks invisible."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn("w:tblPr"), {})
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is None:
        borders = tbl.makeelement(qn("w:tblBorders"), {})
        tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = tbl.makeelement(qn(f"w:{edge}"), {})
            borders.append(el)
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")


def _para_spacing(paragraph, after=None, before=None):
    """Set paragraph spacing."""
    fmt = paragraph.paragraph_format
    if after is not None:
        fmt.space_after = after
    if before is not None:
        fmt.space_before = before
    fmt.line_spacing = C.LINE_SPACING


def _prep_image(uploaded_file, max_w_inches, max_h_inches):
    """Open an uploaded image, resize proportionally, return BytesIO + dimensions."""
    buf = io.BytesIO(uploaded_file.getvalue())
    img = Image.open(buf)
    img = img.convert("RGB")  # normalise RGBA / palette images

    # Calculate proportional size capped by max dims
    w_px, h_px = img.size
    dpi = 96
    max_w_px = int(max_w_inches * dpi)
    max_h_px = int(max_h_inches * dpi)
    ratio = min(max_w_px / w_px, max_h_px / h_px, 1.0)
    new_w = int(w_px * ratio)
    new_h = int(h_px * ratio)
    img = img.resize((new_w, new_h), Image.LANCZOS)

    out = io.BytesIO()
    img.save(out, format="PNG")
    out.seek(0)

    # Return width in inches for docx (capped)
    width_in = min(new_w / dpi, max_w_inches)
    return out, Inches(width_in)


# ── Document Builder ─────────────────────────────────────────────────────────

def generate_docx(data: dict) -> io.BytesIO:
    """
    Build a .docx that replicates the LaTeX template layout.

    Expected keys in `data`:
        student_name, graduate_program, research_topic, sponsor,
        degree, year, contact_email, advisor, career_goal,
        headshot (UploadedFile),
        abstract_p1, abstract_p2 (str, p2 may be empty),
        figure_1, caption_1, figure_2, caption_2
    """
    doc = Document()

    # ── Page setup ──
    for section in doc.sections:
        section.top_margin = C.PAGE_MARGIN
        section.bottom_margin = C.PAGE_MARGIN
        section.left_margin = C.PAGE_MARGIN
        section.right_margin = C.PAGE_MARGIN

    # ──────────────────────────────────────────────────────────────────────
    # 1. HEADER — headshot | student information
    # ──────────────────────────────────────────────────────────────────────
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(header_table)

    # Column widths: headshot 2" + gutter 0.2" | rest
    col_headshot = header_table.columns[0]
    col_info = header_table.columns[1]
    total_width = Inches(6.5)  # 8.5" page - 2×1" margins
    headshot_col_w = C.HEADSHOT_WIDTH + C.HEADER_GUTTER
    col_headshot.width = headshot_col_w
    col_info.width = total_width - headshot_col_w

    cell_photo = header_table.rows[0].cells[0]
    cell_info = header_table.rows[0].cells[1]

    # -- Headshot --
    if data.get("headshot"):
        img_buf, img_w = _prep_image(data["headshot"], 2.0, 2.0)
        p = cell_photo.paragraphs[0]
        run = p.add_run()
        run.add_picture(img_buf, width=C.HEADSHOT_WIDTH)
    else:
        p = cell_photo.paragraphs[0]
        run = p.add_run("[No Headshot]")
        _apply(run, italic=True)

    # -- Info block --
    p = cell_info.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Student name (large bold)
    run = p.add_run(data["student_name"])
    _apply(run, font_size=C.FONT_SIZE_NAME, bold=True)
    run.add_break()

    # Info lines
    info_lines = [
        ("Research topic: ", data["research_topic"])
    ]
    if data.get("sponsor"):
        info_lines.append(("Sponsor: ", data["sponsor"]))
    
    info_lines.extend([
        ("Degree objective: ", f"{data['degree']} ({data['year']})"),
        ("Contact: ", data["contact_email"]),
        ("Advisor: ", data["advisor"]),
        ("Career goal: ", data["career_goal"]),
    ])

    for label, value in info_lines:
        run_label = p.add_run(label)
        _apply(run_label, bold=True)
        run_value = p.add_run(value)
        _apply(run_value)
        run_value.add_break()

    _para_spacing(p, after=Pt(0))

    # ──────────────────────────────────────────────────────────────────────
    # 2. ABSTRACT
    # ──────────────────────────────────────────────────────────────────────
    spacer = doc.add_paragraph()
    _para_spacing(spacer, before=C.SPACING_BEFORE_SECTION, after=Pt(0))

    heading_p = doc.add_paragraph()
    run = heading_p.add_run("Abstract")
    _apply(run, font_size=C.FONT_SIZE_HEADING, bold=True)
    _para_spacing(heading_p, after=C.SPACING_AFTER)

    # Paragraph 1 (required)
    abs1 = doc.add_paragraph()
    _add_parsed_text(abs1, data["abstract_p1"])
    abs1.paragraph_format.first_line_indent = Inches(0)
    _para_spacing(abs1, after=C.SPACING_AFTER)

    # Paragraph 2 (optional)
    if data.get("abstract_p2", "").strip():
        abs2 = doc.add_paragraph()
        _add_parsed_text(abs2, data["abstract_p2"])
        abs2.paragraph_format.first_line_indent = Inches(0)
        _para_spacing(abs2, after=C.SPACING_AFTER)

    # ──────────────────────────────────────────────────────────────────────
    # 3. FIGURES (side-by-side or single)
    # ──────────────────────────────────────────────────────────────────────
    
    # Per new layout rule: if they have two paragraphs, there shouldn't be figures.
    has_fig1 = data.get("figure_1") is not None and not data.get("abstract_p2", "").strip()
    has_fig2 = data.get("figure_2") is not None and not data.get("abstract_p2", "").strip()

    if has_fig1 or has_fig2:
        spacer = doc.add_paragraph()
        _para_spacing(spacer, before=C.SPACING_BEFORE_SECTION, after=Pt(0))

        if has_fig1 and has_fig2:
            # Two-column table for side-by-side figures
            fig_table = doc.add_table(rows=2, cols=2)  # row 0: images, row 1: captions
            fig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            _remove_table_borders(fig_table)
            half = Inches(3.25)
            fig_table.columns[0].width = half
            fig_table.columns[1].width = half

            for idx, col in enumerate([0, 1]):
                fig_key = f"figure_{idx + 1}"
                cap_key = f"caption_{idx + 1}"

                # Image cell
                img_buf, img_w = _prep_image(data[fig_key], 3.0, 2.2)
                cell = fig_table.rows[0].cells[col]
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(img_buf, width=img_w)

                # Caption cell
                cap_cell = fig_table.rows[1].cells[col]
                cap_p = cap_cell.paragraphs[0]
                cap_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                caption_text = data.get(cap_key, "")
                
                _add_parsed_text(cap_p, f"Figure {idx + 1}. ", font_size=C.FONT_SIZE_CAPTION, italic=True)
                _add_parsed_text(cap_p, caption_text, font_size=C.FONT_SIZE_CAPTION, italic=True)

        else:
            # Single figure — centered
            fig_key = "figure_1" if has_fig1 else "figure_2"
            cap_key = "caption_1" if has_fig1 else "caption_2"
            fig_num = 1 if has_fig1 else 2

            img_buf, img_w = _prep_image(data[fig_key], 4.5, 2.7)
            fig_p = doc.add_paragraph()
            fig_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = fig_p.add_run()
            run.add_picture(img_buf, width=img_w)

            caption_text = data.get(cap_key, "")
            if caption_text:
                cap_p = doc.add_paragraph()
                cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = cap_p.add_run(f"Figure {fig_num}. {caption_text}")
                _apply(run, font_size=C.FONT_SIZE_CAPTION, italic=True)

    # ── Save to memory ──
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
